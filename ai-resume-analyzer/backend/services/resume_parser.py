import re
import io
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz)"""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = ""
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            return text
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def extract_email(text: str) -> Optional[str]:
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None

def extract_phone(text: str) -> Optional[str]:
    pattern = r'(\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None

def extract_linkedin(text: str) -> Optional[str]:
    pattern = r'linkedin\.com/in/([a-zA-Z0-9\-]+)'
    match = re.search(pattern, text, re.IGNORECASE)
    return f"linkedin.com/in/{match.group(1)}" if match else None

def extract_github(text: str) -> Optional[str]:
    pattern = r'github\.com/([a-zA-Z0-9\-]+)'
    match = re.search(pattern, text, re.IGNORECASE)
    return f"github.com/{match.group(1)}" if match else None

def extract_name_simple(text: str) -> Optional[str]:
    """Heuristic: first non-empty line that looks like a name"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
            if not any(char in line for char in ['@', '.com', '|', '/', '+']):
                return line
    return lines[0] if lines else None

def estimate_experience_years(text: str) -> float:
    """Estimate total years of experience from text"""
    years_patterns = [
        r'(\d+)\+?\s*years?\s+of\s+experience',
        r'(\d+)\+?\s*years?\s+experience',
        r'experience\s+of\s+(\d+)\+?\s*years?',
    ]
    for pattern in years_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return float(match.group(1))

    # Try to count date ranges
    date_pattern = r'(20\d{2}|19\d{2})\s*[-–]\s*(20\d{2}|19\d{2}|present|current|now)'
    matches = re.findall(date_pattern, text, re.IGNORECASE)
    total = 0
    import datetime
    current_year = datetime.datetime.now().year
    for start, end in matches:
        try:
            s = int(start)
            e = current_year if end.lower() in ['present', 'current', 'now'] else int(end)
            total += max(0, e - s)
        except:
            pass
    return min(total, 30) if total > 0 else 0

def calculate_ats_score(text: str, skills: list, experience_years: float) -> dict:
    """Calculate ATS score based on multiple factors"""
    score_breakdown = {}

    # Keyword optimization (presence of action verbs, quantified achievements)
    action_verbs = ['led', 'managed', 'developed', 'created', 'implemented', 'designed',
                    'built', 'improved', 'increased', 'reduced', 'achieved', 'delivered',
                    'launched', 'optimized', 'automated', 'analyzed', 'coordinated']
    verb_count = sum(1 for verb in action_verbs if verb.lower() in text.lower())
    keyword_score = min(100, (verb_count / len(action_verbs)) * 100 + 20)

    # Quantified achievements
    numbers = re.findall(r'\d+%|\$\d+|\d+[KMB]', text)
    keyword_score = min(100, keyword_score + len(numbers) * 5)
    score_breakdown['keyword_optimization'] = int(keyword_score)

    # Skills match score
    skills_score = min(100, len(skills) * 5 + 30) if skills else 20
    score_breakdown['skills_match'] = int(skills_score)

    # Formatting score
    formatting_score = 60
    if extract_email(text): formatting_score += 10
    if extract_phone(text): formatting_score += 10
    if len(text) > 500: formatting_score += 10
    if len(text) < 5000: formatting_score += 10  # not too long
    score_breakdown['formatting'] = min(100, formatting_score)

    # Experience relevance
    exp_score = min(100, experience_years * 10 + 40) if experience_years > 0 else 40
    score_breakdown['experience_relevance'] = int(exp_score)

    # Education score
    edu_keywords = ['bachelor', 'master', 'phd', 'degree', 'university', 'college', 'b.s.', 'm.s.', 'b.tech', 'm.tech']
    edu_found = any(kw in text.lower() for kw in edu_keywords)
    edu_score = 85 if edu_found else 50
    score_breakdown['education_score'] = edu_score

    # Overall weighted score
    overall = int(
        score_breakdown['keyword_optimization'] * 0.25 +
        score_breakdown['skills_match'] * 0.30 +
        score_breakdown['formatting'] * 0.15 +
        score_breakdown['experience_relevance'] * 0.20 +
        score_breakdown['education_score'] * 0.10
    )

    return {
        'overall': overall,
        **score_breakdown,
        'breakdown': score_breakdown
    }
